"""docx_export.py — Genera documentos Word (.docx): legajo por caso y export total.

Usa python-docx. Pensado para abogados: un "legajo" es un documento completo de un
caso (datos + agenda + adjuntos con links) listo para imprimir o pasarle a otro
profesional. `exportar_todo` arma un Word con todos los casos (uno por sección).

Funciones públicas:
  legajo_caso(caso, eventos, adjuntos, hoy) -> bytes
  exportar_todo(casos, eventos_por_caso, adjuntos_por_caso, hoy) -> bytes

MIME para st.download_button:
  application/vnd.openxmlformats-officedocument.wordprocessingml.document
"""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import casos_logic as cl

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_AZUL = RGBColor(0x1F, 0x3A, 0x5F)


# ----- helpers de formato -----------------------------------------------------

def _add_hyperlink(paragraph, url: str, text: str):
    """Agrega un link CLICKEABLE (azul subrayado) a un párrafo de python-docx."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _titulo_caso(caso: dict) -> str:
    return (caso.get("cliente") or "(sin cliente)")


def _campo_valor(caso: dict, campo: dict) -> str:
    val = caso.get(campo["clave"], "")
    if val in (None, ""):
        return ""
    if campo["tipo"] == "monto":
        return cl.fmt_money(val)
    return str(val)


def _seccion(doc, texto: str):
    h = doc.add_heading(texto, level=1)
    for run in h.runs:
        run.font.color.rgb = _AZUL


# ----- render de un caso (reutilizado por legajo y export total) --------------

def _render_caso(doc, caso: dict, eventos: list, adjuntos: list):
    # Encabezado del caso
    p = doc.add_paragraph()
    r = p.add_run(_titulo_caso(caso))
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = _AZUL
    doc.add_paragraph(
        f"{caso.get('rubro','')}  ·  {caso.get('area','')}  ·  "
        f"Estado: {caso.get('estado_caso') or 'Abierto'}")
    cid = doc.add_paragraph(f"Identificador interno: {caso.get('id','')}")
    cid.runs[0].font.size = Pt(8)
    cid.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Datos del caso (solo campos del área, no vacíos)
    _seccion(doc, "Datos del caso")
    campos = [c for c in cl.campos_de(caso.get("rubro", ""), caso.get("area", ""))
              if c["clave"] not in ("cliente", "estado_caso")]
    filas = [(c["etiqueta"], _campo_valor(caso, c)) for c in campos]
    filas = [(et, v) for et, v in filas if v]
    if filas:
        tabla = doc.add_table(rows=0, cols=2)
        tabla.style = "Light Grid Accent 1"
        for et, v in filas:
            celdas = tabla.add_row().cells
            celdas[0].text = et
            celdas[1].text = v
            celdas[0].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph("Sin datos cargados.")

    # Agenda
    _seccion(doc, "Agenda — audiencias y plazos")
    evs = cl.proximos_eventos(eventos, date.today(), incluir_cumplidos=True)
    if evs:
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Light Grid Accent 1"
        hdr = tabla.rows[0].cells
        for i, t in enumerate(("Fecha", "Tipo", "Detalle", "Estado")):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True
        for e in evs:
            estado = e.get("estado") or "Pendiente"
            if estado != cl.ESTADO_EVENTO_CUMPLIDO:
                estado = e.get("etiqueta", estado)  # ej. "En 3 día(s)" / "Venció hace..."
            c = tabla.add_row().cells
            c[0].text = e.get("fecha", "")
            c[1].text = e.get("tipo", "")
            c[2].text = e.get("descripcion", "") or "—"
            c[3].text = estado
    else:
        doc.add_paragraph("Sin audiencias ni plazos cargados.")

    # Adjuntos
    _seccion(doc, "Documentación adjunta")
    if adjuntos:
        for a in sorted(adjuntos, key=lambda x: x.get("creado_ts", "")):
            linea = doc.add_paragraph(style="List Bullet")
            cab = f"{a.get('tipo','')}"
            if a.get("fecha"):
                cab += f" ({a['fecha']})"
            desc = a.get("descripcion", "")
            r = linea.add_run(f"{cab}: ")
            r.bold = True
            if desc:
                linea.add_run(f"{desc} — ")
            url = a.get("url", "")
            if url:
                _add_hyperlink(linea, url, "abrir archivo")
            else:
                linea.add_run("(sin archivo)")
    else:
        doc.add_paragraph("Sin archivos adjuntos.")


# ----- API pública ------------------------------------------------------------

def legajo_caso(caso: dict, eventos: list, adjuntos: list, hoy=None) -> bytes:
    """Documento Word completo de UN caso (para pasarle a otro abogado)."""
    doc = Document()
    titulo = doc.add_heading("LEGAJO DEL CASO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _render_caso(doc, caso, eventos, adjuntos)

    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = pie.add_run(f"Generado el {cl.fmt_fecha(hoy or date.today())} · Gestor de Casos")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def exportar_todo(casos: list, eventos_por_caso: dict, adjuntos_por_caso: dict,
                  hoy=None) -> bytes:
    """Documento Word con TODOS los casos (uno por página)."""
    doc = Document()
    titulo = doc.add_heading("CASOS DEL ESTUDIO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph(
        f"{len(casos)} caso(s) · Generado el {cl.fmt_fecha(hoy or date.today())}")
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    for i, caso in enumerate(casos):
        if i > 0:
            doc.add_page_break()
        _render_caso(doc, caso,
                     eventos_por_caso.get(caso.get("id"), []),
                     adjuntos_por_caso.get(caso.get("id"), []))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
